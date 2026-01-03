import customtkinter as ctk
from tkinter import messagebox, filedialog
import threading
import uuid
try:
    from docx import Document
except ImportError:
    Document = None
from .service import DocumentAnalysisService

class DocAnalystFrame(ctk.CTkFrame):
    def __init__(self, master, app):
        super().__init__(master, fg_color="transparent")
        self.app = app
        self.service = DocumentAnalysisService(app.data_manager)
        
        self.documents = [] # List of {"name": filename, "content": text}
        self.chat_history = [] # List of {"role": "...", "content": "..."}
        
        self.current_conversation_id = None
        self.current_title = "Nouvelle conversation"
        
        self.suggestion_window = None
        self.suggestion_frame = None

        self.build_ui()
        # Load history initially
        self.after(100, self.refresh_history_list)

    def build_ui(self):
        # 1. Header
        header = ctk.CTkFrame(self, height=60, corner_radius=0)
        header.pack(fill="x", padx=20, pady=10)
        
        btn_back = ctk.CTkButton(
            header,
            text="< Accueil",
            width=100,
            height=32,
            fg_color=("gray70", "gray30"),
            corner_radius=16,
            command=self.app.show_home,
        )
        btn_back.pack(side="left", padx=10, pady=10)
        
        self.lbl_title = ctk.CTkLabel(header, text="📄 Analyse de Documents (Beta)", font=("Arial", 20, "bold"))
        self.lbl_title.pack(side="left", padx=20)

        # 2. Main Content
        content = ctk.CTkFrame(self, fg_color="transparent")
        content.pack(fill="both", expand=True, padx=20, pady=10)
        
        # New Layout: Sidebar (1) | Chat (3) | Document Controls (1)
        content.grid_columnconfigure(0, weight=1) # Sidebar
        content.grid_columnconfigure(1, weight=4) # Chat area
        content.grid_columnconfigure(2, weight=1) # Controls
        content.grid_rowconfigure(0, weight=1)

        # --- Left: Sidebar (History) ---
        sidebar_panel = ctk.CTkFrame(content, width=200)
        sidebar_panel.grid(row=0, column=0, sticky="nsew", padx=(0, 10))
        
        ctk.CTkLabel(sidebar_panel, text="Historique", font=("Arial", 14, "bold")).pack(pady=10)
        
        self.btn_new_chat = ctk.CTkButton(
            sidebar_panel,
            text="+ Nouvelle conv.",
            fg_color="green",
            command=self.start_new_conversation
        )
        self.btn_new_chat.pack(pady=(0, 10), padx=10, fill="x")
        
        self.scroll_history = ctk.CTkScrollableFrame(sidebar_panel, fg_color="transparent")
        self.scroll_history.pack(fill="both", expand=True, padx=5, pady=5)


        # --- Middle: Chat Area ---
        chat_panel = ctk.CTkFrame(content)
        chat_panel.grid(row=0, column=1, sticky="nsew", padx=(0, 10))
        chat_panel.grid_rowconfigure(0, weight=1)
        chat_panel.grid_columnconfigure(0, weight=1)

        self.chat_display = ctk.CTkTextbox(chat_panel, state="disabled", font=("Arial", 12))
        self.chat_display.grid(row=0, column=0, sticky="nsew", padx=10, pady=10)

        input_frame = ctk.CTkFrame(chat_panel, height=50)
        input_frame.grid(row=1, column=0, sticky="ew", padx=10, pady=10)
        
        self.msg_entry = ctk.CTkEntry(input_frame, placeholder_text="Posez votre question sur le document...")
        self.msg_entry.pack(side="left", fill="x", expand=True, padx=(0, 10))
        self.msg_entry.bind("<Return>", self.send_message)
        self.msg_entry.bind("<KeyRelease>", self.check_mentions)
        self.msg_entry.bind("<FocusOut>", lambda e: self.after(200, self.hide_suggestions))
        
        self.btn_send = ctk.CTkButton(input_frame, text="Envoyer", width=100, command=self.send_message)
        self.btn_send.pack(side="right")
        
        # Loading Indicator (Progress Bar Style)
        self.progress_bar = ctk.CTkProgressBar(
            chat_panel,
            mode="indeterminate",
            height=8,
            corner_radius=4,
            progress_color=("#4CAF50", "#4CAF50"),
            fg_color=("gray85", "gray25")
        )
        # Initially hidden
        self.progress_bar.set(0)

        # --- Right: Controls ---
        controls_panel = ctk.CTkFrame(content, width=250)
        controls_panel.grid(row=0, column=2, sticky="nsew")
        
        ctk.CTkLabel(controls_panel, text="Document", font=("Arial", 14, "bold")).pack(pady=20)
        
        self.btn_upload = ctk.CTkButton(
            controls_panel, 
            text="📂 Charger Documents (PDF/TXT)",
            command=self.upload_document
        )
        self.btn_upload.pack(pady=10, padx=20, fill="x")

        self.btn_clear = ctk.CTkButton(
            controls_panel,
            text="🗑️ Vider la liste",
            fg_color="firebrick",
            hover_color="darkred",
            command=self.clear_documents
        )
        self.btn_clear.pack(pady=5, padx=20, fill="x")
        
        
        self.lbl_filename = ctk.CTkLabel(controls_panel, text="Aucun fichier chargé", text_color="gray")
        self.lbl_filename.pack(pady=5)
        
        ctk.CTkLabel(controls_panel, text="Paramètres", font=("Arial", 14, "bold")).pack(pady=(30, 10))
        
        # Provider Selection
        self.var_provider = ctk.StringVar(value="")
        self.combo_provider = ctk.CTkOptionMenu(controls_panel, variable=self.var_provider, values=[])
        self.combo_provider.pack(pady=10, padx=20, fill="x")
        self.update_provider_list()

        # Export Button at Bottom
        self.btn_export = ctk.CTkButton(
            controls_panel, 
            text="💾  Exporter l'analyse (Word)",
            font=("Arial", 14, "bold"),
            height=40,
            fg_color="#2B579A", # Professional Word Blue
            hover_color="#1E3E70",
            command=self.export_to_word
        )
        self.btn_export.pack(side="bottom", pady=20, padx=20, fill="x")

        
    def update_provider_list(self):
        # Hardcoded Hugging Face models for Document Analysis
        self.hf_models = [
            "IAKA (Interne)",
            "Qwen 2.5 72B (Hugging Face)"
        ]
        
        self.combo_provider.configure(values=self.hf_models)
        
        settings = self.app.data_manager.get_settings()
        default_provider = settings.get("doc_analyst_provider")
        
        if default_provider and default_provider in self.hf_models:
            self.var_provider.set(default_provider)
        else:
            self.var_provider.set(self.hf_models[0])
            
            
    # --- Conversation Management ---
    
    def start_new_conversation(self):
        self.current_conversation_id = None
        self.current_title = "Nouvelle conversation"
        self.documents = []
        self.chat_history = []
        
        self.update_document_label()
        
        self.chat_display.configure(state="normal")
        self.chat_display.delete("1.0", "end")
        self.chat_display.configure(state="disabled")
        
        self.lbl_title.configure(text="📄 Analyse de Documents (Nouv. conv.)")
        self.append_chat("System", "Nouvelle conversation démarrée.")

    def refresh_history_list(self):
        # Clear sidebar
        for widget in self.scroll_history.winfo_children():
            widget.destroy()
            
        conversations = self.service.get_all_conversations()
        # Sort by updated_at desc usually
        try:
            conversations.sort(key=lambda x: x.get('updated_at', ''), reverse=True)
        except:
            pass # Handle legacy or malformed dates gracefully

        for conv in conversations:
            f = ctk.CTkFrame(self.scroll_history, fg_color="transparent")
            f.pack(fill="x", pady=2)
            
            # Determine highlighting
            is_active = (self.current_conversation_id == conv["id"])
            bg_color = ("gray85", "gray25") if is_active else "transparent"
            text_weight = "bold" if is_active else "normal"

            # Title Button
            btn = ctk.CTkButton(
                f,
                text=conv.get("title", "Sans titre"),
                anchor="w",
                fg_color=bg_color,
                text_color=("black", "white"),
                hover_color=("gray80", "gray30"),
                font=("Arial", 12, text_weight),
                height=30,
                command=lambda c=conv: self.load_conversation(c)
            )
            btn.pack(side="left", fill="x", expand=True)

            # Rename Button
            rename_btn = ctk.CTkButton(
                f,
                text="✎",
                width=30,
                fg_color="transparent",
                text_color="gray",
                hover_color=("gray90", "gray40"),
                command=lambda cid=conv["id"]: self.rename_conversation_ui(cid)
            )
            rename_btn.pack(side="right", padx=(2, 0))
            
            # Delete Button
            del_btn = ctk.CTkButton(
                f,
                text="X",
                width=30,
                fg_color="transparent",
                text_color="red",
                hover_color=("mistyrose", "darkred"),
                command=lambda cid=conv["id"]: self.delete_conversation_ui(cid)
            )
            del_btn.pack(side="right")

    def rename_conversation_ui(self, conversation_id):
        dialog = ctk.CTkInputDialog(text="Nouveau nom de conversation :", title="Renommer")
        new_title = dialog.get_input()
        if new_title:
            self.service.rename_conversation(conversation_id, new_title)
            
            # Update current title if we are renaming the active conversation
            if self.current_conversation_id == conversation_id:
                self.current_title = new_title
                
            self.refresh_history_list()

    def load_conversation(self, conversation):
        self.current_conversation_id = conversation["id"]
        self.current_title = conversation.get("title", "Sans titre")
        self.documents = conversation.get("documents", [])
        self.chat_history = conversation.get("messages", [])
        
        # Restore UI
        self.lbl_title.configure(text=f"📄 {self.current_title}")
        self.update_document_label()
        
        self.chat_display.configure(state="normal")
        self.chat_display.delete("1.0", "end")
        
        # Replay history
        for msg in self.chat_history:
            sender = "Vous" if msg["role"] == "user" else "Assistant"
            self.chat_display.insert("end", f"\n[{sender}]: {msg['content']}\n")
            
        self.chat_display.see("end")
        self.chat_display.configure(state="disabled")

    def delete_conversation_ui(self, conversation_id):
        if messagebox.askyesno("Confirmer", "Supprimer cette conversation ?"):
            self.service.delete_conversation(conversation_id)
            if self.current_conversation_id == conversation_id:
                self.start_new_conversation()
            self.refresh_history_list()

    def _save_current_state(self):
        # Generate ID if new
        if not self.current_conversation_id:
            self.current_conversation_id = str(uuid.uuid4())
            
        new_title = self.service.save_conversation(
            self.current_conversation_id,
            self.current_title,
            self.chat_history,
            self.documents
        )
        
        if new_title != self.current_title:
            self.current_title = new_title
            self.lbl_title.configure(text=f"📄 {self.current_title}")
            
        self.refresh_history_list()


    def upload_document(self):
        file_paths = filedialog.askopenfilenames(
            filetypes=[("Documents", "*.pdf *.txt")]
        )
        if not file_paths:
            return

        loaded_count = 0
        errors = []

        for file_path in file_paths:
            success, result = self.service.extract_text(file_path)
            if success:
                filename = file_path.split("/")[-1]
                self.documents.append({"name": filename, "content": result})
                loaded_count += 1
            else:
                errors.append(f"{file_path.split('/')[-1]}: {result}")

        self.update_document_label()
        
        if loaded_count > 0:
            # Collect names of newly added documents (last 'loaded_count' items)
            new_docs = [d['name'] for d in self.documents[-loaded_count:]]
            msg = f"Document(s) ajouté(s) : {', '.join(new_docs)}"
            if errors:
                msg += f"\nErreurs non chargées: {', '.join(errors)}"
            self.append_chat("System", msg)
            
            # Save state after upload (always save to create/update session)
            self._save_current_state()
            
        elif errors:
            messagebox.showerror("Erreurs", "\n".join(errors))

    def clear_documents(self):
        if messagebox.askyesno("Confirmation", "Cela effacera les documents de la conversation ACTUELLE."):
            self.documents = []
            # We don't necessarily clear chat history, just docs? 
            # User said "Vider la liste", implying docs. But usually context depends on docs.
            # I will keep history but warn.
            self.update_document_label()
            self.append_chat("System", "Liste des documents vidée.")
            self._save_current_state()

    def update_document_label(self):
        count = len(self.documents)
        if count == 0:
            self.lbl_filename.configure(text="Aucun fichier chargé", text_color="gray")
        else:
            total_chars = sum(len(d['content']) for d in self.documents)
            
            # Context Limit Warnings (Approximation: 1 token ~= 4 chars)
            if total_chars > 120000: # ~30k tokens
                color = "red"
                icon = "⛔ ATTENTION: Très volumineux"
            elif total_chars > 32000: # ~8k tokens (standard context)
                color = "orange"
                icon = "⚠️ Volumineux"
            else:
                color = "green"
                icon = "✅"
                
            self.lbl_filename.configure(
                text=f"{icon} {count} document(s)\n(~{total_chars} caractères)", 
                text_color=color
            )

    def send_message(self, event=None):
        if not self.documents:
            messagebox.showwarning("Info", "Veuillez d'abord charger un document.")
            return
            
        msg = self.msg_entry.get().strip()
        if not msg:
            return

        provider = self.var_provider.get()
        if not provider or provider == "Aucun provider configuré":
            messagebox.showerror("Erreur", "Veuillez sélectionner un provider LLM.")
            return

        self.append_chat("Vous", msg)
        self.msg_entry.delete(0, "end")
        self.btn_send.configure(state="disabled")
        
        # Show loading indicator
        self.progress_bar.grid(row=2, column=0, sticky="ew", padx=20, pady=(0, 10))
        self.progress_bar.start()
        
        # Save before sending (async)
        # self._save_current_state() # wait for response to save complete pair?

        # Threaded call
        thread = threading.Thread(target=self._chat_thread, args=(msg, provider))
        thread.daemon = True
        thread.start()

    def _chat_thread(self, msg, provider):
        # Combine all documents
        full_context = ""
        for doc in self.documents:
            full_context += f"--- DEBUT DOCUMENT: {doc['name']} ---\n"
            full_context += doc['content'] + "\n"
            full_context += f"--- FIN DOCUMENT: {doc['name']} ---\n\n"

        success, response = self.service.chat_with_document(
            full_context,
            msg,
            self.chat_history,
            provider
        )
        self.after(0, lambda: self._on_response(success, response, msg))

    def _on_response(self, success, response, user_msg):
        self.progress_bar.stop()
        self.progress_bar.grid_forget() # Hide loading
        self.btn_send.configure(state="normal")
        if success:
            self.append_chat("Assistant", response)
            # Add to history
            self.chat_history.append({"role": "user", "content": user_msg})
            self.chat_history.append({"role": "assistant", "content": response})
            
            # Save state
            self._save_current_state()
            
        else:
            self.append_chat("System", f"Erreur: {response}")

    def append_chat(self, sender, text):
        self.chat_display.configure(state="normal")
        self.chat_display.insert("end", f"\n[{sender}]: {text}\n")
        self.chat_display.see("end")
        self.chat_display.configure(state="disabled")

    def export_to_word(self):
        if not self.chat_history:
            messagebox.showinfo("Info", "Aucune conversation à exporter.")
            return

        if Document is None:
             messagebox.showerror("Erreur", "La librairie 'python-docx' n'est pas installée.")
             return

        file_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Documents", "*.docx")]
        )
        if not file_path:
            return

        try:
            doc = Document()
            # Clean filename from label text
            # Clean filename from label text
            raw_filename = self.lbl_filename.cget("text")
            # Simple fallback for name
            clean_filename = f"Export_{len(self.documents)}_Docs"
            
            doc.add_heading(f"Analyse de Document : {clean_filename}", 0)

            for msg in self.chat_history:
                sender = "Utilisateur" if msg["role"] == "user" else "Assistant"
                p = doc.add_paragraph()
                p.add_run(f"[{sender}]: ").bold = True
                p.add_run(msg["content"])
                
            doc.save(file_path)
            messagebox.showinfo("Succès", "Export réussi !")
        except Exception as e:
            messagebox.showerror("Erreur", f"Erreur lors de l'export : {e}")

    # --- Autocomplete Logic ---
    def check_mentions(self, event):
        # Don't trigger on internal navigation keys
        if event.keysym in ("Up", "Down", "Left", "Right", "Return", "Escape"):
             if event.keysym == "Escape":
                 self.hide_suggestions()
             return

        text = self.msg_entry.get()
        cursor_pos = self.msg_entry._entry.index("insert")
        
        # Look backwards from cursor for '@'
        # Simple implementation: detect if we are currently typing a word starting with @
        # Get text up to cursor
        current_text = text[:cursor_pos]
        
        # Find last @
        last_at = current_text.rfind("@")
        
        if last_at != -1:
            # Check if there are spaces between @ and cursor (allow spaces in filenames? usually no for mentions but files have spaces)
            # Let's say mention ends at space for now, or ensure we are "in" the mention
            # Actually, standard behavior is @trigger until space.
            query = current_text[last_at+1:]
            
            # If query contains space, abort (unless we want to support spaces in mentions, which is hard)
            # For simplicity, we abort on space to avoid false positives
            if " " in query:
                self.hide_suggestions()
                return

            self.show_suggestions(query)
        else:
            self.hide_suggestions()

    def show_suggestions(self, query):
        # Filter documents
        matches = [d["name"] for d in self.documents if query.lower() in d["name"].lower()]
        
        if not matches:
            self.hide_suggestions()
            return

        # Create window if not exists
        if self.suggestion_window is None or not self.suggestion_window.winfo_exists():
            self.suggestion_window = ctk.CTkToplevel(self)
            self.suggestion_window.wm_overrideredirect(True)
            self.suggestion_window.wm_attributes("-topmost", True)
            # transparent not supported on Toplevel directly in this context without specific system calls
            # We just let it have default background which will be covered by suggestion_frame
            
            # Position logic
            entry_x = self.msg_entry.winfo_rootx()
            entry_y = self.msg_entry.winfo_rooty()
            entry_h = self.msg_entry.winfo_height()
            
            # Offset slightly
            self.suggestion_window.geometry(f"+{entry_x + 10}+{entry_y - (len(matches) * 35) - 10}")

            self.suggestion_frame = ctk.CTkFrame(self.suggestion_window, corner_radius=10, fg_color=("gray90", "gray20"), border_width=1, border_color="gray50")
            self.suggestion_frame.pack(fill="both", expand=True)

        # Update content
        # Clear old
        for widget in self.suggestion_frame.winfo_children():
            widget.destroy()

        # Add buttons
        for doc_name in matches:
             btn = ctk.CTkButton(
                 self.suggestion_frame, 
                 text=doc_name, 
                 anchor="w",
                 fg_color="transparent", 
                 text_color=("black", "white"),
                 hover_color=("gray80", "gray30"),
                 height=30,
                 command=lambda n=doc_name: self.insert_mention(n, query)
             )
             btn.pack(fill="x", padx=5, pady=2)
             
        # Resize window
        total_h = len(matches) * 34 + 10
        width = 250
        
        entry_x = self.msg_entry.winfo_rootx()
        entry_y = self.msg_entry.winfo_rooty()
        # appear above
        self.suggestion_window.geometry(f"{width}x{total_h}+{entry_x}+{entry_y - total_h - 5}")

    def hide_suggestions(self):
        if self.suggestion_window:
            self.suggestion_window.destroy()
            self.suggestion_window = None
            self.suggestion_frame = None

    def insert_mention(self, filename, query):
        current_text = self.msg_entry.get()
        cursor_pos = self.msg_entry._entry.index("insert")
        
        # We need to replace the last occurance of '@query' before cursor
        # Re-locate start index
        start_pos = cursor_pos - len(query) - 1 # -1 for @
        
        # Construct new text
        prefix = current_text[:start_pos]
        suffix = current_text[cursor_pos:]
        
        new_text = prefix + f"@{filename} " + suffix
        
        self.msg_entry.delete(0, "end")
        self.msg_entry.insert(0, new_text)
        
        # Correct cursor placement is hard without index math, but appending at end is safe
        # Or calculate new pos
        new_cursor = start_pos + len(filename) + 2 # +1 for @, +1 for space
        self.msg_entry._entry.icursor(new_cursor)
        
        self.hide_suggestions()
